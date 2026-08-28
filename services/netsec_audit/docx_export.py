# -*- coding: utf-8 -*-
"""Generatore di report Word (.docx) per NetSec Audit."""

import io
from typing import Any, Dict
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls


def _set_cell_background(cell, hex_color: str):
    """Imposta il colore di sfondo di una cella della tabella."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color.lstrip("#")}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def _set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Imposta margini interni di una cella in dxa."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def generate_audit_docx(data: Dict[str, Any]) -> bytes:
    """Genera un documento .docx completo e formattato dal payload di audit."""
    doc = Document()

    # Impostazioni pagina A4 portrait con margini standard
    for section in doc.sections:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    # Palette colori
    C_NAVY = RGBColor(15, 23, 42)
    C_SLATE = RGBColor(71, 85, 105)
    C_MUTED = RGBColor(100, 116, 139)
    C_DARK = RGBColor(30, 41, 59)
    C_RED = RGBColor(220, 38, 38)
    C_GREEN = RGBColor(22, 163, 74)
    C_AMBER = RGBColor(217, 119, 6)

    lang = data.get("lang", "it")
    is_en = (lang == "en")

    device = data.get("device_name") or data.get("device_ip") or "Device"
    benchmark = data.get("benchmark_title") or data.get("benchmark") or "CIS Benchmark"
    generated = data.get("generated") or ""
    vendor = data.get("vendor", "")
    platform_name = "Cisco IOS XE" if vendor == "ios" else "FortiOS" if vendor == "fortios" else "Linux" if vendor == "linux" else vendor.upper()
    score = data.get("score")
    summary = data.get("summary", {})
    rules = data.get("rules", [])

    # Header documento
    p_title = doc.add_paragraph()
    run_brand = p_title.add_run("SentinelNet  |  ")
    run_brand.bold = True
    run_brand.font.size = Pt(18)
    run_brand.font.color.rgb = C_NAVY
    run_brand.font.name = "Arial"

    run_sub = p_title.add_run("Security Compliance Report" if is_en else "Report di Conformità di Sicurezza")
    run_sub.bold = True
    run_sub.font.size = Pt(16)
    run_sub.font.color.rgb = C_SLATE
    run_sub.font.name = "Arial"
    p_title.paragraph_format.space_after = Pt(2)

    p_bm = doc.add_paragraph()
    run_bm = p_bm.add_run(benchmark)
    run_bm.font.size = Pt(11)
    run_bm.font.color.rgb = C_SLATE
    run_bm.font.name = "Arial"
    p_bm.paragraph_format.space_after = Pt(10)

    # Tabella Metadati
    t_meta = doc.add_table(rows=2, cols=4)
    t_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_headers = ["TARGET DEVICE" if is_en else "DISPOSITIVO",
                    "PLATFORM" if is_en else "PIATTAFORMA",
                    "BENCHMARK",
                    "GENERATED ON" if is_en else "DATA GENERAZIONE"]
    meta_values = [str(device), platform_name, benchmark, str(generated)]

    for i in range(4):
        c_hdr = t_meta.cell(0, i)
        c_hdr.text = meta_headers[i]
        c_hdr.paragraphs[0].runs[0].font.size = Pt(8.5)
        c_hdr.paragraphs[0].runs[0].font.bold = True
        c_hdr.paragraphs[0].runs[0].font.color.rgb = C_MUTED
        c_hdr.paragraphs[0].runs[0].font.name = "Arial"
        _set_cell_background(c_hdr, "f8fafc")
        _set_cell_margins(c_hdr, top=60, bottom=40, left=100, right=100)

        c_val = t_meta.cell(1, i)
        c_val.text = meta_values[i]
        c_val.paragraphs[0].runs[0].font.size = Pt(9.5)
        c_val.paragraphs[0].runs[0].font.bold = True
        c_val.paragraphs[0].runs[0].font.color.rgb = C_NAVY
        c_val.paragraphs[0].runs[0].font.name = "Arial"
        _set_cell_background(c_val, "f8fafc")
        _set_cell_margins(c_val, top=40, bottom=80, left=100, right=100)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Banner valutazione parziale se presente
    unknown_count = summary.get("unknown", 0)
    total_count = summary.get("total", len(rules))
    if unknown_count > 0:
        p_warn = doc.add_paragraph()
        p_warn.paragraph_format.space_before = Pt(4)
        p_warn.paragraph_format.space_after = Pt(8)
        run_w_title = p_warn.add_run("Valutazione parziale: " if not is_en else "Partial assessment: ")
        run_w_title.bold = True
        run_w_title.font.size = Pt(9.5)
        run_w_title.font.color.rgb = C_AMBER
        run_w_title.font.name = "Arial"

        assessed = total_count - unknown_count
        w_text = (f"{assessed} controlli su {total_count} sono stati valutati; {unknown_count} non valutabili per assenza sezioni nel file."
                  if not is_en else f"{assessed} of {total_count} checks assessed; {unknown_count} not assessable due to missing sections.")
        run_w_body = p_warn.add_run(w_text)
        run_w_body.font.size = Pt(9.5)
        run_w_body.font.color.rgb = C_DARK
        run_w_body.font.name = "Arial"

    # Tabella KPI Cards
    t_kpi = doc.add_table(rows=2, cols=5)
    t_kpi.alignment = WD_TABLE_ALIGNMENT.CENTER
    kpi_cols = [
        (f"{score}%" if score is not None else "—", "COMPLIANCE SCORE" if is_en else "PUNTEGGIO", "10b981" if (score or 0) >= 80 else "f59e0b" if (score or 0) >= 50 else "ef4444"),
        (str(summary.get("passed", 0)), "COMPLIANT (PASS)" if is_en else "CONFORMI (PASS)", "10b981"),
        (str(summary.get("failed", 0)), "NON-COMPLIANT (FAIL)" if is_en else "NON CONFORMI (FAIL)", "ef4444"),
        (str(summary.get("warned", 0)), "WARNINGS (WARN)" if is_en else "AVVISI (WARN)", "f59e0b"),
        (str(summary.get("unknown", 0)), "NOT ASSESSABLE" if is_en else "NON VALUTABILI", "64748b"),
    ]

    for i, (val, label, hex_col) in enumerate(kpi_cols):
        c_val = t_kpi.cell(0, i)
        c_val.text = val
        r_v = c_val.paragraphs[0].runs[0]
        r_v.font.size = Pt(16)
        r_v.font.bold = True
        r_v.font.name = "Arial"
        c_val.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_background(c_val, "ffffff")
        _set_cell_margins(c_val, top=80, bottom=20, left=60, right=60)

        c_lbl = t_kpi.cell(1, i)
        c_lbl.text = label
        r_l = c_lbl.paragraphs[0].runs[0]
        r_l.font.size = Pt(7.5)
        r_l.font.bold = True
        r_l.font.color.rgb = C_MUTED
        r_l.font.name = "Arial"
        c_lbl.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_background(c_lbl, "ffffff")
        _set_cell_margins(c_lbl, top=20, bottom=80, left=60, right=60)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Tabella Risultati Controlli (Findings Table)
    t_find = doc.add_table(rows=1, cols=5)
    t_find.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Intestazioni tabella
    headers = [
        ("ID / REF", Inches(1.1)),
        ("CHECK, GUIDANCE & EVIDENCE" if is_en else "CONTROLLO, GUIDA ED EVIDENZE", Inches(3.7)),
        ("SEVERITY" if is_en else "SEVERITÀ", Inches(0.7)),
        ("RESULT" if is_en else "ESITO", Inches(0.7)),
        ("REMEDIATION (CLI)" if is_en else "RIMEDIO (CLI)", Inches(1.5)),
    ]

    hdr_row = t_find.rows[0]
    for i, (h_text, width) in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.width = width
        cell.text = h_text
        r = cell.paragraphs[0].runs[0]
        r.font.size = Pt(8.5)
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.name = "Arial"
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (2, 3) else WD_ALIGN_PARAGRAPH.LEFT
        _set_cell_background(cell, "0f172a")
        _set_cell_margins(cell, top=80, bottom=80, left=80, right=80)

    # Righe controlli
    for r in rules:
        row = t_find.add_row()
        for idx, (_, width) in enumerate(headers):
            row.cells[idx].width = width
            _set_cell_margins(row.cells[idx], top=60, bottom=60, left=80, right=80)

        # Colonna 0: ID / REF
        c0 = row.cells[0]
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_after = Pt(2)
        r_id = p0.add_run(r.get("id", ""))
        r_id.font.bold = True
        r_id.font.size = Pt(8.5)
        r_id.font.name = "Consolas"
        r_id.font.color.rgb = C_NAVY

        ref = r.get("ref")
        lvl = r.get("level")
        if ref or lvl:
            p_ref = c0.add_paragraph()
            p_ref.paragraph_format.space_after = Pt(0)
            if ref:
                r_ref = p_ref.add_run(f"[{ref}] ")
                r_ref.font.size = Pt(7.5)
                r_ref.font.name = "Consolas"
                r_ref.font.color.rgb = C_MUTED
            if lvl:
                r_lvl = p_ref.add_run(f"L{lvl}")
                r_lvl.font.size = Pt(7.5)
                r_lvl.font.name = "Consolas"
                r_lvl.font.color.rgb = C_MUTED

        # Colonna 1: CHECK, GUIDANCE & EVIDENCE
        c1 = row.cells[1]
        p1_title = c1.paragraphs[0]
        p1_title.paragraph_format.space_after = Pt(2)
        r_title = p1_title.add_run(r.get("title", ""))
        r_title.font.bold = True
        r_title.font.size = Pt(9.5)
        r_title.font.name = "Arial"
        r_title.font.color.rgb = C_NAVY

        if r.get("detail"):
            p1_desc = c1.add_paragraph()
            p1_desc.paragraph_format.space_after = Pt(4)
            r_desc = p1_desc.add_run(r.get("detail", ""))
            r_desc.font.size = Pt(8.5)
            r_desc.font.name = "Arial"
            r_desc.font.color.rgb = C_DARK

        # Guidance (Why / Impact / Default)
        g = r.get("guidance") or {}
        if g.get("why") or g.get("impact") or g.get("default"):
            p1_g = c1.add_paragraph()
            p1_g.paragraph_format.space_after = Pt(4)
            if g.get("why"):
                r_why_lbl = p1_g.add_run("Perché conta: " if not is_en else "Why it matters: ")
                r_why_lbl.font.bold = True
                r_why_lbl.font.size = Pt(8.5)
                r_why_lbl.font.name = "Arial"
                r_why_lbl.font.color.rgb = C_NAVY
                r_why_val = p1_g.add_run(f"{g.get('why')}\n")
                r_why_val.font.size = Pt(8.5)
                r_why_val.font.name = "Arial"
                r_why_val.font.color.rgb = C_DARK

            if g.get("impact"):
                r_imp_lbl = p1_g.add_run("Impatto del fix: " if not is_en else "Impact of the fix: ")
                r_imp_lbl.font.bold = True
                r_imp_lbl.font.size = Pt(8.5)
                r_imp_lbl.font.name = "Arial"
                r_imp_lbl.font.color.rgb = C_NAVY
                r_imp_val = p1_g.add_run(f"{g.get('impact')}\n")
                r_imp_val.font.size = Pt(8.5)
                r_imp_val.font.name = "Arial"
                r_imp_val.font.color.rgb = C_DARK

            if g.get("default"):
                r_def_lbl = p1_g.add_run("Default di fabbrica: " if not is_en else "Factory default: ")
                r_def_lbl.font.bold = True
                r_def_lbl.font.size = Pt(8.5)
                r_def_lbl.font.name = "Arial"
                r_def_lbl.font.color.rgb = C_NAVY
                r_def_val = p1_g.add_run(f"{g.get('default')}")
                r_def_val.font.size = Pt(8.5)
                r_def_val.font.name = "Arial"
                r_def_val.font.color.rgb = C_DARK

        # Verifica su apparato
        if r.get("audit"):
            p1_aud = c1.add_paragraph()
            p1_aud.paragraph_format.space_after = Pt(4)
            r_aud_lbl = p1_aud.add_run("Comando di verifica: " if not is_en else "Verify command: ")
            r_aud_lbl.font.bold = True
            r_aud_lbl.font.size = Pt(8)
            r_aud_lbl.font.name = "Arial"
            r_aud_lbl.font.color.rgb = C_SLATE
            r_aud_cmd = p1_aud.add_run(r.get("audit"))
            r_aud_cmd.font.size = Pt(8)
            r_aud_cmd.font.name = "Consolas"
            r_aud_cmd.font.color.rgb = C_NAVY

        # Evidenze
        ev_list = r.get("evidence") or []
        if ev_list:
            p1_ev = c1.add_paragraph()
            p1_ev.paragraph_format.space_after = Pt(2)
            r_ev_hdr = p1_ev.add_run("EVIDENZE NELLA CONFIGURAZIONE:\n" if not is_en else "EVIDENCE IN CONFIGURATION:\n")
            r_ev_hdr.font.bold = True
            r_ev_hdr.font.size = Pt(7.5)
            r_ev_hdr.font.name = "Arial"
            r_ev_hdr.font.color.rgb = C_RED

            for ev in ev_list:
                p_item = c1.add_paragraph()
                p_item.paragraph_format.space_after = Pt(1)
                line_str = f"Riga {ev.get('line')}" if ev.get("line") else "—"
                raw_ctx = ev.get("context", "")
                ctx_clean = raw_ctx.replace("firewall policy / ", "Policy ID #").replace("policy / ", "Policy ID #")

                r_ev_l = p_item.add_run(f"[{line_str}] ")
                r_ev_l.font.size = Pt(7.5)
                r_ev_l.font.name = "Consolas"
                r_ev_l.font.color.rgb = C_MUTED

                if ctx_clean:
                    r_ev_c = p_item.add_run(f"{ctx_clean}: ")
                    r_ev_c.font.bold = True
                    r_ev_c.font.size = Pt(8)
                    r_ev_c.font.name = "Consolas"
                    r_ev_c.font.color.rgb = C_DARK

                r_ev_t = p_item.add_run(str(ev.get("text", "")))
                r_ev_t.font.bold = True
                r_ev_t.font.size = Pt(8)
                r_ev_t.font.name = "Consolas"
                r_ev_t.font.color.rgb = C_RED

        # Colonna 2: SEVERITY
        c2 = row.cells[2]
        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sev = str(r.get("severity", "MEDIUM")).upper()
        r_sev = p2.add_run(sev)
        r_sev.font.bold = True
        r_sev.font.size = Pt(8)
        r_sev.font.name = "Arial"
        r_sev.font.color.rgb = C_RED if sev in ("CRITICAL", "HIGH") else C_AMBER if sev == "MEDIUM" else C_MUTED

        # Colonna 3: RESULT
        c3 = row.cells[3]
        p3 = c3.paragraphs[0]
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        st = str(r.get("status", "UNKNOWN")).upper()
        r_st = p3.add_run(st)
        r_st.font.bold = True
        r_st.font.size = Pt(8)
        r_st.font.name = "Arial"
        r_st.font.color.rgb = C_GREEN if st == "PASS" else C_RED if st == "FAIL" else C_AMBER if st == "WARN" else C_MUTED

        # Colonna 4: REMEDIATION (CLI)
        c4 = row.cells[4]
        p4 = c4.paragraphs[0]
        rem = r.get("remediation") or "—"
        r_rem = p4.add_run(rem)
        r_rem.font.size = Pt(8)
        r_rem.font.name = "Consolas"
        r_rem.font.color.rgb = RGBColor(2, 132, 199)
        _set_cell_background(c4, "f0f9ff")

    # Footer
    p_foot = doc.add_paragraph()
    p_foot.paragraph_format.space_before = Pt(16)
    r_foot = p_foot.add_run("SentinelNet Security Audit Engine  •  Generato automaticamente da configurazione apparato")
    r_foot.font.size = Pt(8)
    r_foot.font.color.rgb = C_MUTED
    r_foot.font.name = "Arial"

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
